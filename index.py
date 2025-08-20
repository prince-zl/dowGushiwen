import requests
from bs4 import BeautifulSoup
import time
import os
import json
import re
from urllib.parse import urljoin, urlparse

# Word文档相关导入
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
    print("✅ python-docx 已安装，将自动生成Word文档")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx 未安装，请运行: pip install python-docx")
    print("   将只生成TXT和JSON文件")

class UniversalNovelCrawler:
    def __init__(self, catalog_url=None):
        self.catalog_url = catalog_url
        self.base_domain = self.get_base_domain(catalog_url) if catalog_url else ""
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            'Connection': 'keep-alive',
            'Referer': self.base_domain,
        }
        self.session = requests.Session()
        self.session.headers.update(self.headers)
        self.book_info = {'title': '未知书名', 'author': '未知作者'}
        
    def get_base_domain(self, url):
        """获取URL的基础域名"""
        if not url:
            return ""
        parsed = urlparse(url)
        return f"{parsed.scheme}://{parsed.netloc}"
    
    def get_book_info_from_catalog(self, soup):
        """从目录页面提取书名和作者信息"""
        try:
            # 常见的书名选择器
            title_selectors = [
                'h1', 'h2', '.book-title', '.title', '#title',
                '.book-name', '.bookname', '.novel-title',
                '[class*="title"]', '[id*="title"]',
                '.main h1', '.content h1', '.header h1'
            ]
            
            book_title = "未知书名"
            for selector in title_selectors:
                title_elem = soup.select_one(selector)
                if title_elem:
                    title_text = title_elem.get_text(strip=True)
                    if title_text and len(title_text) > 0 and len(title_text) < 50:
                        book_title = title_text
                        break
            
            # 常见的作者选择器
            author_selectors = [
                '.author', '.book-author', '#author',
                '[class*="author"]', '[id*="author"]',
                '.writer', '.novelist'
            ]
            
            book_author = "未知作者"
            for selector in author_selectors:
                author_elem = soup.select_one(selector)
                if author_elem:
                    author_text = author_elem.get_text(strip=True)
                    # 清理作者文本
                    author_text = re.sub(r'作者[：:]\s*', '', author_text)
                    author_text = re.sub(r'著[：:]\s*', '', author_text)
                    if author_text and len(author_text) > 0 and len(author_text) < 30:
                        book_author = author_text
                        break
            
            # 如果在标题中找到作者信息
            if "作者" in book_title:
                parts = re.split(r'[作者：:]', book_title)
                if len(parts) >= 2:
                    book_title = parts[0].strip()
                    book_author = parts[1].strip()
            
            self.book_info = {
                'title': book_title,
                'author': book_author
            }
            
            print(f"📚 检测到书籍信息:")
            print(f"   书名: {book_title}")
            print(f"   作者: {book_author}")
            
        except Exception as e:
            print(f"⚠️  获取书籍信息失败: {e}")
            self.book_info = {'title': '未知书名', 'author': '未知作者'}
    
    def parse_chapter_list(self):
        """从目录页面解析章节列表"""
        if not self.catalog_url:
            print("❌ 未设置目录URL")
            return []
        
        print(f"🔍 正在解析目录页面: {self.catalog_url}")
        
        try:
            response = self.session.get(self.catalog_url, timeout=15)
            response.encoding = 'utf-8'
            
            if response.status_code != 200:
                print(f"❌ 目录页面访问失败: {response.status_code}")
                return []
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 提取书籍信息
            self.get_book_info_from_catalog(soup)
            
            chapters = []
            
            # 策略1: 寻找包含章节链接的容器
            container_selectors = [
                '.chapter-list', '.catalogue', '.catalog', '.list',
                '.content-list', '.book-list', '#list', '#catalog',
                '.main .list', '.content .list', '[class*="chapter"]',
                '[class*="catalog"]', '[class*="list"]', 'ul', 'ol',
                '.directory', '.index', '.toc'
            ]
            
            chapter_container = None
            for selector in container_selectors:
                container = soup.select_one(selector)
                if container:
                    # 检查容器内是否有足够多的链接
                    links = container.find_all('a', href=True)
                    if len(links) >= 3:  # 至少3个链接才认为是章节容器
                        chapter_container = container
                        print(f"✅ 找到章节容器: {selector} (包含 {len(links)} 个链接)")
                        break
            
            if chapter_container:
                links = chapter_container.find_all('a', href=True)
                
                for link in links:
                    href = link.get('href')
                    title = link.get_text(strip=True)
                    
                    # 过滤掉明显不是章节的链接
                    if (title and len(title) > 0 and len(title) < 100 and
                        not any(skip in title.lower() for skip in 
                               ['首页', '书架', '登录', '注册', '搜索', '排行', 
                                'home', 'login', 'register', 'search', 'rank',
                                '上一页', '下一页', '返回', '目录'])):
                        
                        # 将相对URL转换为绝对URL
                        full_url = urljoin(self.catalog_url, href)
                        
                        chapters.append({
                            'title': title,
                            'url': full_url
                        })
            
            # 策略2: 如果没找到容器，直接在整个页面查找链接
            if not chapters:
                print("🔧 策略1失败，尝试在整个页面查找章节链接...")
                
                all_links = soup.find_all('a', href=True)
                
                # 分析链接模式，找出可能的章节链接
                potential_chapters = []
                
                for link in all_links:
                    href = link.get('href')
                    title = link.get_text(strip=True)
                    
                    # 检查链接文本是否像章节标题
                    if (title and len(title) > 3 and len(title) < 100 and
                        (re.search(r'第[一二三四五六七八九十百千万\d]+[章回节卷集部]', title) or
                         re.search(r'[第]?\d+[章回节卷集部]', title) or
                         re.search(r'chapter\s*\d+', title, re.I) or
                         '章' in title or '回' in title or '节' in title)):
                        
                        full_url = urljoin(self.catalog_url, href)
                        potential_chapters.append({
                            'title': title,
                            'url': full_url
                        })
                
                # 如果找到了疑似章节的链接
                if potential_chapters:
                    chapters = potential_chapters
                    print(f"✅ 通过模式匹配找到 {len(chapters)} 个疑似章节")
                else:
                    # 最后的尝试：找到最多链接的区域
                    print("🔧 尝试最后策略：分析链接密度...")
                    
                    # 找到包含最多链接的div或section
                    containers = soup.find_all(['div', 'section', 'ul', 'ol'])
                    best_container = None
                    max_links = 0
                    
                    for container in containers:
                        links_in_container = container.find_all('a', href=True)
                        if len(links_in_container) > max_links and len(links_in_container) >= 5:
                            max_links = len(links_in_container)
                            best_container = container
                    
                    if best_container:
                        links = best_container.find_all('a', href=True)
                        for link in links:
                            href = link.get('href')
                            title = link.get_text(strip=True)
                            
                            if title and len(title) > 0 and len(title) < 100:
                                full_url = urljoin(self.catalog_url, href)
                                chapters.append({
                                    'title': title,
                                    'url': full_url
                                })
                        
                        print(f"✅ 从最佳容器找到 {len(chapters)} 个链接")
            
            # 去重
            unique_chapters = []
            seen_urls = set()
            for chapter in chapters:
                if chapter['url'] not in seen_urls:
                    unique_chapters.append(chapter)
                    seen_urls.add(chapter['url'])
            
            print(f"📋 总共解析到 {len(unique_chapters)} 个唯一章节")
            
            # 🎯 重要：对章节进行排序，确保从第一章开始
            if unique_chapters:
                sorted_chapters = self.sort_chapters(unique_chapters)
                print(f"📊 章节排序完成，识别到 {len(sorted_chapters)} 个有序章节")
                unique_chapters = sorted_chapters
            
            # 显示前几个章节作为预览
            if unique_chapters:
                print("📖 章节预览:")
                for i, chapter in enumerate(unique_chapters[:5], 1):
                    print(f"   {i}. {chapter['title']}")
                if len(unique_chapters) > 5:
                    print(f"   ... 还有 {len(unique_chapters) - 5} 个章节")
            
            return unique_chapters
            
        except Exception as e:
            print(f"❌ 解析目录页面失败: {e}")
            return []
    
    def get_chapter_content(self, chapter_url, chapter_title):
        """获取单个章节内容 - 严格按照p标签分段，并智能合并标题"""
        try:
            print(f"  📖 获取: {chapter_title}")
            
            response = self.session.get(chapter_url, timeout=15)
            response.encoding = 'utf-8'
            
            if response.status_code != 200:
                print(f"    ❌ HTTP错误: {response.status_code}")
                return {"title": chapter_title, "content": ""}
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 移除不需要的元素
            for element in soup(['script', 'style', 'nav', 'header', 'footer', 'aside']):
                element.decompose()
            
            content = ""
            
            # 策略1: 寻找主要内容容器
            content_selectors = [
                '.content', '.main-content', '.chapter-content', '.text-content',
                '#content', '#main-content', '#chapter-content', '#text-content',
                '.post-content', '.entry-content', '.article-content',
                '.main .content', '.container .content', '[class*="content"]',
                '.main3 .left .cont', '.main3 .cont', '.cont',  # 古诗文网特有
                '.chapter', '.article', '.post', '.entry'
            ]
            
            for selector in content_selectors:
                content_container = soup.select_one(selector)
                if content_container:
                    print(f"    ✅ 找到内容容器: {selector}")
                    
                    # 严格提取所有p标签作为段落
                    paragraphs = content_container.find_all('p')
                    
                    if paragraphs:
                        paragraph_texts = []
                        for p in paragraphs:
                            p_text = p.get_text(strip=True)
                            if p_text and len(p_text) > 5:  # 过滤太短的段落
                                paragraph_texts.append(p_text)
                        
                        if paragraph_texts:
                            content = '\n\n'.join(paragraph_texts)  # 每个p标签间用双换行分隔
                            print(f"    ✅ 严格按p标签提取到 {len(paragraph_texts)} 个段落")
                            break
                    
                    # 如果容器内没有p标签，检查其他可能的段落结构
                    if not content:
                        # 检查div或span作为段落
                        div_paragraphs = content_container.find_all(['div', 'span'])
                        if len(div_paragraphs) > 1:
                            para_texts = []
                            for div in div_paragraphs:
                                div_text = div.get_text(strip=True)
                                if div_text and len(div_text) > 10:
                                    para_texts.append(div_text)
                            
                            if len(para_texts) > 1:
                                content = '\n\n'.join(para_texts)
                                print(f"    ✅ 按div/span标签提取到 {len(para_texts)} 个段落")
                                break
                        
                        # 检查br标签分割的内容
                        container_html = str(content_container)
                        if '<br' in container_html.lower():
                            br_separated = re.sub(r'<br[^>]*?/?>', '\n||PARAGRAPH_BREAK||\n', container_html)
                            clean_text = BeautifulSoup(br_separated, 'html.parser').get_text()
                            paragraphs = [p.strip() for p in clean_text.split('||PARAGRAPH_BREAK||')]
                            paragraphs = [p for p in paragraphs if p and len(p) > 10]
                            
                            if paragraphs:
                                content = '\n\n'.join(paragraphs)
                                print(f"    ✅ 按br标签分段提取到 {len(paragraphs)} 个段落")
                                break
            
            # 策略2: 如果容器策略失败，直接在整个页面中查找所有p标签
            if not content or len(content) < 100:
                print(f"    🔧 策略1失败，在整个页面查找p标签...")
                
                all_paragraphs = soup.find_all('p')
                
                if all_paragraphs:
                    paragraph_texts = []
                    for p in all_paragraphs:
                        p_text = p.get_text(strip=True)
                        # 过滤掉明显的导航、菜单、版权信息
                        if (p_text and len(p_text) > 15 and 
                            not any(skip in p_text.lower() for skip in 
                                   ['导航', '菜单', '登录', '注册', '首页', '版权', 'copyright', 
                                    '关于我们', '联系我们', '用户协议', '隐私政策', '意见反馈',
                                    '上一章', '下一章', '返回目录', '书签', '收藏'])):
                            paragraph_texts.append(p_text)
                    
                    if len(paragraph_texts) > 2:
                        content = '\n\n'.join(paragraph_texts)
                        print(f"    ✅ 从全页面严格按p标签提取到 {len(paragraph_texts)} 个段落")
            
            # 最终内容验证和格式化
            if content:
                # 清理多余的空行，但保持双换行的段落分隔
                content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)
                content = content.strip()
                
                # 🎯 智能标题合并处理
                merged_title, cleaned_content = self.merge_titles(chapter_title, content)
                
                # 验证内容质量
                paragraph_count = cleaned_content.count('\n\n') + 1
                
                if len(cleaned_content) > 50 and paragraph_count >= 1:
                    print(f"    ✅ 最终成功 ({len(cleaned_content)} 字符, {paragraph_count} 个段落)")
                    if merged_title != chapter_title:
                        print(f"    🔗 标题已合并: {merged_title}")
                    return {"title": merged_title, "content": cleaned_content}
                else:
                    print(f"    ⚠️  内容质量不足 ({len(cleaned_content)} 字符, {paragraph_count} 个段落)")
                    return {"title": chapter_title, "content": cleaned_content}  # 即使质量不足也返回
            
            print(f"    ❌ 所有策略都未能提取到有效的分段内容")
            return {"title": chapter_title, "content": ""}
            
        except Exception as e:
            print(f"    ❌ 获取失败: {e}")
            return {"title": chapter_title, "content": ""}
    
    def merge_titles(self, catalog_title, content):
        """智能合并目录标题和内容标题"""
        if not content:
            return catalog_title, content
        
        # 获取内容的第一行 - 尝试多种分割方式
        lines_double = content.split('\n\n')  # 双换行分割
        lines_single = content.split('\n')    # 单换行分割
        
        # 选择第一行
        first_line = ""
        remaining_content = content
        
        # 优先使用双换行分割，如果第一段太长则使用单换行
        if lines_double and lines_double[0].strip():
            first_line = lines_double[0].strip()
            if len(first_line) <= 150:  # 如果第一段不太长，认为可能是标题
                remaining_content = '\n\n'.join(lines_double[1:]).strip() if len(lines_double) > 1 else ""
            else:
                # 第一段太长，尝试单换行分割
                if lines_single and lines_single[0].strip():
                    first_line = lines_single[0].strip()
                    remaining_content = '\n'.join(lines_single[1:]).strip() if len(lines_single) > 1 else ""
        elif lines_single and lines_single[0].strip():
            first_line = lines_single[0].strip()
            remaining_content = '\n'.join(lines_single[1:]).strip() if len(lines_single) > 1 else ""
        
        if not first_line:
            return catalog_title, content
        
        print(f"    🔍 检查第一行: {first_line[:50]}{'...' if len(first_line) > 50 else ''}")
        
        # 判断第一行是否像标题
        is_title = self.is_likely_title(first_line)
        
        if is_title:
            # 合并标题：目录标题 + 内容标题
            merged_title = self.combine_titles(catalog_title, first_line)
            
            print(f"    🔗 检测到内容标题，已合并: {first_line}")
            return merged_title, remaining_content
        else:
            # 第一行不是标题，保持原样
            print(f"    ❌ 第一行不被识别为标题，保持原样")
            return catalog_title, content
    
    def is_likely_title(self, text):
        """判断文本是否可能是标题"""
        if not text:
            return False
        
        # 太长的文本不太可能是标题
        if len(text) > 200:
            return False
        
        # 太短的文本也不太可能是完整标题
        if len(text) < 3:
            return False
        
        print(f"    🔍 标题判断分析: '{text}'")
        
        # 强标题特征 - 如果匹配则几乎确定是标题
        strong_title_patterns = [
            r'第[一二三四五六七八九十百千万\d]+[章回节卷集部篇]',  # 第X章
            r'^[第]?\d+[章回节卷集部篇]',  # X章 或 第X章
            r'chapter\s*\d+',  # Chapter X
            r'^[\d\s\-\.]+[章回节卷集部篇]',  # 数字开头+章节词
            r'^第.*[章回节卷集部篇]',  # 以"第"开头，以章节词结尾
        ]
        
        for pattern in strong_title_patterns:
            if re.search(pattern, text, re.I):
                print(f"    ✅ 匹配强标题模式: {pattern}")
                return True
        
        # 中等标题特征
        medium_indicators = 0
        
        # 包含章节关键词
        chapter_keywords = ['章', '回', '节', '卷', '集', '部', '篇', 'chapter']
        if any(keyword in text.lower() for keyword in chapter_keywords):
            medium_indicators += 2
            print(f"    📝 包含章节关键词 (+2)")
        
        # 包含序号
        if re.search(r'\d+', text):
            medium_indicators += 1
            print(f"    🔢 包含数字 (+1)")
        
        # 不以句号结尾（正文通常以句号结尾）
        if not text.endswith(('。', '！', '？', '.', '!', '?')):
            medium_indicators += 1
            print(f"    📄 不以句号结尾 (+1)")
        
        # 长度适中
        if 5 <= len(text) <= 50:
            medium_indicators += 1
            print(f"    📏 长度适中 (+1)")
        
        # 不包含过多标点符号（正文通常标点较多）
        punctuation_ratio = len(re.findall(r'[，。！？、；：""''（）【】《》]', text)) / len(text)
        if punctuation_ratio < 0.3:
            medium_indicators += 1
            print(f"    📝 标点符号少 (+1)")
        
        # 包含常见标题词汇
        title_words = ['初', '始', '末', '终', '新', '老', '大', '小', '上', '下', '前', '后', 
                      '东', '西', '南', '北', '入', '出', '来', '去', '见', '遇', '战', '斗']
        if any(word in text for word in title_words):
            medium_indicators += 1
            print(f"    📚 包含标题常用词 (+1)")
        
        print(f"    📊 总分: {medium_indicators}/7")
        
        # 如果累积指标足够高，认为是标题
        if medium_indicators >= 3:
            print(f"    ✅ 根据综合指标判断为标题")
            return True
        
        print(f"    ❌ 综合指标不足，判断为非标题")
        return False
    
    def combine_titles(self, catalog_title, content_title):
        """合并目录标题和内容标题"""
        # 清理标题
        catalog_clean = catalog_title.strip()
        content_clean = content_title.strip()
        
        print(f"    🔗 合并标题:")
        print(f"       目录: {catalog_clean}")
        print(f"       内容: {content_clean}")
        
        # 如果内容标题就是目录标题的一部分，直接返回目录标题
        if content_clean in catalog_clean:
            print(f"    ✅ 内容标题包含在目录标题中，使用目录标题")
            return catalog_clean
        
        if catalog_clean in content_clean:
            print(f"    ✅ 目录标题包含在内容标题中，使用内容标题")
            return content_clean
        
        # 检查是否有相同的章节号
        catalog_chapter_num = self.extract_chapter_number(catalog_clean)
        content_chapter_num = self.extract_chapter_number(content_clean)
        
        print(f"    🔢 章节号 - 目录: {catalog_chapter_num}, 内容: {content_chapter_num}")
        
        if catalog_chapter_num and content_chapter_num and catalog_chapter_num == content_chapter_num:
            # 如果章节号相同，智能合并
            catalog_without_num = re.sub(r'第[一二三四五六七八九十百千万\d]+[章回节卷集部篇]\s*', '', catalog_clean)
            content_without_num = re.sub(r'第[一二三四五六七八九十百千万\d]+[章回节卷集部篇]\s*', '', content_clean)
            content_without_num = re.sub(r'^\d+[章回节卷集部篇]\s*', '', content_without_num)
            
            print(f"    📝 去除章节号后 - 目录: '{catalog_without_num}', 内容: '{content_without_num}'")
            
            if catalog_without_num and content_without_num and catalog_without_num != content_without_num:
                result = f"{catalog_clean} {content_without_num}"
                print(f"    ✅ 合并结果: {result}")
                return result
            else:
                print(f"    ✅ 章节号相同且内容重复，使用目录标题")
                return catalog_clean
        
        # 检查是否有明显的重复词汇
        catalog_words = set(catalog_clean.replace('第', '').replace('章', '').replace('回', '').split())
        content_words = set(content_clean.replace('第', '').replace('章', '').replace('回', '').split())
        
        common_words = catalog_words & content_words
        if len(common_words) > 0 and len(common_words) >= len(catalog_words) * 0.5:
            print(f"    ⚠️  发现大量重复词汇: {common_words}，使用较长的标题")
            return catalog_clean if len(catalog_clean) >= len(content_clean) else content_clean
        
        # 默认合并：目录标题 + 内容标题
        result = f"{catalog_clean} {content_clean}"
        print(f"    ✅ 默认合并结果: {result}")
        return result
    
    def extract_chapter_number(self, title):
        """提取章节号"""
        # 匹配各种章节号格式
        patterns = [
            r'第([一二三四五六七八九十百千万]+)[章回节卷集部篇]',
            r'第(\d+)[章回节卷集部篇]',
            r'^(\d+)[章回节卷集部篇]',
            r'[第]?(\d+)[章回节卷集部篇]',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, title)
            if match:
                return match.group(1)
        
        return None
    
    def sort_chapters(self, chapters):
        """对章节进行排序，确保正确的阅读顺序"""
        print("🔄 正在对章节进行排序...")
        
        def extract_chapter_sort_key(chapter_title):
            """提取章节的排序关键字"""
            title = chapter_title.lower().strip()
            
            # 尝试提取各种格式的章节号
            patterns = [
                # 中文数字
                (r'第([一二三四五六七八九十百千万]+)[章回节卷集部篇]', chinese_to_number),
                # 阿拉伯数字
                (r'第(\d+)[章回节卷集部篇]', int),
                (r'^(\d+)[章回节卷集部篇]', int),
                (r'[章回节卷集部篇](\d+)', int),
                # Chapter格式
                (r'chapter\s*(\d+)', int),
                # 纯数字开头
                (r'^(\d+)', int),
            ]
            
            for pattern, converter in patterns:
                match = re.search(pattern, title)
                if match:
                    try:
                        num = converter(match.group(1))
                        return (0, num)  # 0表示是正常章节，num是章节号
                    except:
                        continue
            
            # 特殊处理一些关键词
            special_keywords = {
                '序章': (-2, 0),
                '序言': (-2, 0), 
                '序': (-2, 0),
                '前言': (-2, 0),
                '楔子': (-1, 0),
                '引子': (-1, 0),
                '开篇': (-1, 0),
                '终章': (999, 999),
                '尾声': (999, 999),
                '后记': (999, 999),
                '番外': (1000, 0),
            }
            
            for keyword, sort_key in special_keywords.items():
                if keyword in title:
                    return sort_key
            
            # 无法识别章节号的，放在最后
            return (500, 999999)
        
        # 中文数字转换函数
        def chinese_to_number(chinese_str):
            """将中文数字转换为阿拉伯数字"""
            chinese_dict = {
                '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
                '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
                '十一': 11, '十二': 12, '十三': 13, '十四': 14, '十五': 15,
                '十六': 16, '十七': 17, '十八': 18, '十九': 19, '二十': 20,
                '二十一': 21, '二十二': 22, '二十三': 23, '二十四': 24, '二十五': 25,
                '二十六': 26, '二十七': 27, '二十八': 28, '二十九': 29, '三十': 30,
                '三十一': 31, '三十二': 32, '三十三': 33, '三十四': 34, '三十五': 35,
                '三十六': 36, '三十七': 37, '三十八': 38, '三十九': 39, '四十': 40,
                '四十一': 41, '四十二': 42, '四十三': 43, '四十四': 44, '四十五': 45,
                '四十六': 46, '四十七': 47, '四十八': 48, '四十九': 49, '五十': 50,
                '五十一': 51, '五十二': 52, '五十三': 53, '五十四': 54, '五十五': 55,
                '五十六': 56, '五十七': 57, '五十八': 58, '五十九': 59, '六十': 60,
                '六十一': 61, '六十二': 62, '六十三': 63, '六十四': 64, '六十五': 65,
                '六十六': 66, '六十七': 67, '六十八': 68, '六十九': 69, '七十': 70,
                '七十一': 71, '七十二': 72, '七十三': 73, '七十四': 74, '七十五': 75,
                '七十六': 76, '七十七': 77, '七十八': 78, '七十九': 79, '八十': 80,
                '八十一': 81, '八十二': 82, '八十三': 83, '八十四': 84, '八十五': 85,
                '八十六': 86, '八十七': 87, '八十八': 88, '八十九': 89, '九十': 90,
                '九十一': 91, '九十二': 92, '九十三': 93, '九十四': 94, '九十五': 95,
                '九十六': 96, '九十七': 97, '九十八': 98, '九十九': 99, '一百': 100,
            }
            
            # 处理更复杂的中文数字
            if chinese_str in chinese_dict:
                return chinese_dict[chinese_str]
            
            # 处理百位数
            if '百' in chinese_str:
                parts = chinese_str.split('百')
                if len(parts) == 2:
                    hundred_part = parts[0] if parts[0] else '一'
                    remainder_part = parts[1]
                    
                    hundred_num = chinese_dict.get(hundred_part, 1) * 100
                    remainder_num = chinese_dict.get(remainder_part, 0) if remainder_part else 0
                    
                    return hundred_num + remainder_num
            
            return 0
        
        # 给每个章节添加排序key
        chapters_with_key = []
        for chapter in chapters:
            sort_key = extract_chapter_sort_key(chapter['title'])
            chapters_with_key.append((sort_key, chapter))
            print(f"   📋 {chapter['title']} → 排序键: {sort_key}")
        
        # 按排序key排序
        chapters_with_key.sort(key=lambda x: x[0])
        
        # 返回排序后的章节列表
        sorted_chapters = [chapter for _, chapter in chapters_with_key]
        
        print(f"✅ 排序完成，顺序:")
        for i, chapter in enumerate(sorted_chapters[:10], 1):
            print(f"   {i}. {chapter['title']}")
        if len(sorted_chapters) > 10:
            print(f"   ... 还有 {len(sorted_chapters) - 10} 个章节")
        
        return sorted_chapters
    
    def crawl_book(self, delay=3, test_mode=False):
        """爬取整本书"""
        if not self.catalog_url:
            print("❌ 未设置目录URL，无法爬取")
            return
        
        mode_text = "测试模式（前3章）" if test_mode else "完整模式（所有章节）"
        print(f"🚀 开始爬取《{self.book_info['title']}》 - {mode_text}")
        print("=" * 60)
        
        # 解析章节列表
        chapters = self.parse_chapter_list()
        
        if not chapters:
            print("❌ 无法获取任何章节信息")
            return
        
        # 测试模式：只爬取前几章
        if test_mode:
            chapters = chapters[:3]
            print(f"🧪 测试模式：只爬取前 {len(chapters)} 章")
        else:
            print(f"🚀 完整模式：准备爬取所有 {len(chapters)} 个章节")
        
        print(f"📚 找到 {len(chapters)} 个章节，开始爬取...")
        print("-" * 60)
        
        # 爬取章节内容
        chapters_data = []
        success_count = 0
        
        for i, chapter in enumerate(chapters, 1):
            print(f"[{i:3d}/{len(chapters)}] {chapter['title']}")
            
            result = self.get_chapter_content(chapter['url'], chapter['title'])
            merged_title = result["title"]
            content = result["content"]
            
            # 计算段落数量
            paragraph_count = content.count('\n\n') + 1 if content else 0
            
            chapter_data = {
                'original_title': chapter['title'],  # 保存原始目录标题
                'title': merged_title,               # 使用合并后的标题
                'url': chapter['url'],
                'content': content,
                'char_count': len(content),
                'paragraph_count': paragraph_count,
                'success': len(content) > 50 and paragraph_count >= 1
            }
            
            chapters_data.append(chapter_data)
            
            if chapter_data['success']:
                success_count += 1
                if merged_title != chapter['title']:
                    print(f"           ✅ 成功 ({paragraph_count} 个段落) - 标题已合并")
                else:
                    print(f"           ✅ 成功 ({paragraph_count} 个段落)")
            else:
                print(f"           ❌ 失败或内容不完整")
            
            # 添加延迟，避免请求过快
            if i < len(chapters):  # 最后一章不需要延迟
                print(f"           ⏳ 等待 {delay} 秒...")
                time.sleep(delay)
        
        # 保存结果
        print("-" * 60)
        print(f"📋 爬取完成:")
        print(f"   总章节: {len(chapters_data)}")
        print(f"   成功: {success_count}")
        print(f"   失败: {len(chapters_data) - success_count}")
        
        # 统计段落信息
        total_paragraphs = sum(ch.get('paragraph_count', 0) for ch in chapters_data)
        total_chars = sum(ch.get('char_count', 0) for ch in chapters_data)
        print(f"   总段落数: {total_paragraphs}")
        print(f"   总字符数: {total_chars:,}")
        
        if success_count > 0:
            # 🎯 创建以书名命名的文件夹
            safe_title = re.sub(r'[<>:"/\\|?*]', '_', self.book_info['title'])
            safe_title = safe_title.replace('  ', ' ').strip('_').strip()
            
            # 确保文件夹名不为空
            if not safe_title or safe_title in ['_', '.', '..']:
                safe_title = f"小说_{int(time.time())}"
            
            # 创建文件夹
            try:
                os.makedirs(safe_title, exist_ok=True)
                print(f"📁 创建文件夹: {safe_title}")
            except Exception as e:
                print(f"⚠️  创建文件夹失败: {e}")
                safe_title = "."  # 保存到当前目录
            
            # 生成文件路径
            txt_path = os.path.join(safe_title, f"{safe_title}.txt")
            json_path = os.path.join(safe_title, f"{safe_title}.json")
            docx_path = os.path.join(safe_title, f"{safe_title}.docx")
            
            # 保存所有格式
            self.save_to_file(chapters_data, txt_path)
            self.save_to_json(chapters_data, json_path)
            
            # 生成Word文档
            if DOCX_AVAILABLE:
                self.save_to_word(chapters_data, docx_path)
            
            print(f"\n🎉 爬取完成！获得 {success_count} 个有效章节，共 {total_paragraphs} 个段落")
            print(f"📁 所有文件已保存到文件夹: {safe_title}")
            
            return safe_title  # 返回文件夹名
        else:
            print("❌ 没有成功获取任何章节内容")
            print("\n💡 建议:")
            print("   1. 检查目录URL是否正确")
            print("   2. 检查网络连接")
            print("   3. 稍后重试（可能遇到频率限制）")
            print("   4. 尝试使用VPN或更换IP")
    
    def save_to_file(self, chapters_data, filename):
        """保存内容到文件 - 保持段落格式"""
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(f"{self.book_info['title']}\n")
                f.write("=" * 60 + "\n")
                f.write(f"作者：{self.book_info['author']}\n")
                f.write("爬取时间：" + time.strftime('%Y-%m-%d %H:%M:%S') + "\n")
                f.write(f"来源：{self.catalog_url}\n")
                f.write("说明：严格按照原网站p标签分段保存，智能合并标题\n")
                f.write("=" * 60 + "\n\n")
                
                total_chars = 0
                total_paragraphs = 0
                
                for i, chapter in enumerate(chapters_data, 1):
                    f.write(f"{chapter['title']}\n")
                    f.write("-" * 50 + "\n\n")
                    
                    if chapter['content']:
                        f.write(chapter['content'])  # 内容已经按段落格式化，直接写入
                        total_chars += len(chapter['content'])
                        total_paragraphs += chapter.get('paragraph_count', 0)
                    else:
                        f.write("[此章节内容获取失败]")
                    
                    f.write(f"\n\n\n")
                
                f.write(f"\n\n总计: {len(chapters_data)} 章, {total_chars:,} 字, {total_paragraphs} 段落\n")
            
            print(f"✅ TXT文件已保存: {filename}")
            
        except Exception as e:
            print(f"❌ 保存TXT文件失败: {e}")
    
    def save_to_json(self, chapters_data, filename):
        """保存为JSON格式"""
        try:
            data = {
                'title': self.book_info['title'],
                'author': self.book_info['author'],
                'crawl_time': time.strftime('%Y-%m-%d %H:%M:%S'),
                'source_url': self.catalog_url,
                'format_note': '严格按照原网站p标签分段，智能合并标题',
                'total_chapters': len(chapters_data),
                'success_chapters': len([ch for ch in chapters_data if ch.get('success', False)]),
                'total_chars': sum(ch['char_count'] for ch in chapters_data),
                'total_paragraphs': sum(ch.get('paragraph_count', 0) for ch in chapters_data),
                'chapters': chapters_data
            }
            
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            
            print(f"✅ JSON文件已保存: {filename}")
            
        except Exception as e:
            print(f"❌ 保存JSON失败: {e}")
    
    def save_to_word(self, chapters_data, filename):
        """保存为Word文档 - 仿照用户提供的排版格式"""
        try:
            print("📝 正在生成Word文档...")
            
            # 创建新的Word文档
            document = Document()
            
            # 设置中文字体 - 全局样式
            document.styles['Normal'].font.name = '楷体'
            document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            document.styles['Normal'].font.size = Pt(14)  # 四号字
            
            # 📝 文档标题 - 楷体
            title = document.add_heading(self.book_info['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.name = '楷体'
            title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            title.runs[0].font.size = Pt(18)  # 小二号
            title.runs[0].bold = True
            
            # 空行
            document.add_paragraph()
            
            # 📝 作者信息 - 楷体
            author_para = document.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f'作者：{self.book_info["author"]}')
            author_run.font.name = '楷体'
            author_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            author_run.font.size = Pt(14)  # 四号字
            
            # 🎯 作者信息后插入分页符
            document.add_page_break()
            
            # 📝 添加章节内容 - 仿照用户排版
            success_count = 0
            
            for i, chapter in enumerate(chapters_data, 1):
                if not chapter.get('success', False) or not chapter.get('content'):
                    continue
                    
                success_count += 1
                print(f"📄 正在处理第 {success_count} 章: {chapter['title']}")
                
                # 🎯 章回标题 - 使用标题1样式，楷体小二号
                title_heading = document.add_heading(chapter['title'], 1)  # 标题1样式
                title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER        # 居中对齐
                title_heading.runs[0].font.name = '楷体'
                title_heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                title_heading.runs[0].font.size = Pt(18)                   # 小二号 = 18pt
                title_heading.runs[0].bold = True
                
                # 调整标题后间距
                title_heading.paragraph_format.space_after = Pt(6)         # 减少标题后间距
                
                # 📝 章节内容 - 仿照用户的段落格式
                content = chapter['content']
                paragraphs = content.split('\n\n')  # 按双换行分段
                
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        # 🎯 创建段落 - 首行缩进两个字符
                        para = document.add_paragraph()
                        para_run = para.add_run(para_text)
                        para_run.font.name = '楷体'
                        para_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                        para_run.font.size = Pt(12)
                        
                        # 🎯 段落格式 - 首行缩进两个字符
                        para.paragraph_format.first_line_indent = Pt(24)      # 首行缩进两个字符 (12pt * 2 = 24pt)
                        para.paragraph_format.left_indent = Inches(0)         # 左对齐
                        para.paragraph_format.space_after = Pt(8)             # 段后间距
                        para.paragraph_format.line_spacing = 1.15             # 行距
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT              # 左对齐
                
                # 章节结束（不添加额外空行，直接分页）
                
                # 🎯 每一回结束后添加分页符（除了最后一回）
                if success_count < len([ch for ch in chapters_data if ch.get('success', False)]):
                    document.add_page_break()
            
            # 保存Word文档
            document.save(filename)
            print(f"✅ Word文档已保存: {filename}")
            print(f"📊 成功处理 {success_count} 个章节")
            print(f"🎨 排版特点：楷体字体、标题1样式小二号、正文四号字、紧凑间距、每回分页")
            
        except Exception as e:
            print(f"❌ 保存Word文档失败: {e}")
            print("💡 可能是python-docx库问题，请检查安装：pip install python-docx")

def main():
    """主函数 - 包含完整的用户选择界面"""
    print("=" * 60)
    print("        通用小说爬虫 - 动态URL版本 v3.0")
    print("        (支持任意小说网站目录页面)")
    print("=" * 60)
    print("🎯 支持大部分小说网站的目录页面")
    print("📝 自动识别书名、作者和章节列表")
    print("🔗 智能合并目录标题和内容标题")
    print("📄 爬取完成后自动生成格式化Word文档")
    print("⚠️  请遵守网站使用条款，仅用于学习研究")
    print("=" * 60)
    
    # 获取用户输入的目录URL
    while True:
        print("\n请输入小说目录页面的URL:")
        print("💡 示例:")
        print("   https://www.gushiwen.cn/guwen/book_ce3ab505d8e6.aspx")
        print("   https://www.某小说网站.com/book/12345/")
        print("   https://m.某网站.com/novel/目录页面")
        print()
        
        catalog_url = input("📝 目录URL: ").strip()
        
        if not catalog_url:
            print("❌ URL不能为空，请重新输入")
            continue
        
        # 简单验证URL格式
        if not catalog_url.startswith(('http://', 'https://')):
            print("❌ 请输入完整的URL（包含http://或https://）")
            continue
        
        # 测试URL是否可访问
        print(f"🔍 正在测试URL访问性: {catalog_url}")
        
        try:
            test_headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            test_response = requests.get(catalog_url, headers=test_headers, timeout=10)
            
            if test_response.status_code == 200:
                print("✅ URL可以正常访问")
                break
            else:
                print(f"⚠️  URL返回状态码 {test_response.status_code}，是否继续？")
                continue_choice = input("继续使用此URL吗？(y/n): ").strip().lower()
                if continue_choice in ['y', 'yes', '是']:
                    break
                else:
                    continue
        except Exception as e:
            print(f"⚠️  URL访问测试失败: {e}")
            continue_choice = input("仍要使用此URL吗？(y/n): ").strip().lower()
            if continue_choice in ['y', 'yes', '是']:
                break
            else:
                continue
    
    # 创建爬虫实例
    crawler = UniversalNovelCrawler(catalog_url)
    
    # 询问用户想要的模式
    while True:
        print("\n请选择爬取模式:")
        print("1. 🧪 测试模式（只爬取前3章，快速验证效果）")
        print("2. 🚀 完整模式（爬取所有章节，生成完整Word文档）")
        print("3. 📊 先测试再决定（推荐）")
        print("4. ❌ 退出程序")
        
        try:
            choice = input("\n请输入选择 (1/2/3/4): ").strip()
            
            if choice == "1":
                print("\n" + "="*50)
                print("🧪 已选择：测试模式")
                print("📝 将爬取前3章并生成测试Word文档...")
                crawler.crawl_book(delay=2, test_mode=True)
                break
                
            elif choice == "2":
                print("\n" + "="*50)
                print("🚀 已选择：完整模式")
                print("📚 将爬取所有章节并生成完整Word文档...")
                
                # 二次确认
                confirm = input("⚠️  完整爬取可能需要较长时间，确认继续？(y/n): ").strip().lower()
                if confirm in ['y', 'yes', '是', '确认']:
                    crawler.crawl_book(delay=3, test_mode=False)
                    break
                else:
                    print("❌ 已取消完整爬取，返回主菜单")
                    continue
                
            elif choice == "3":
                print("\n" + "="*50)
                print("📊 推荐模式：先测试再决定")
                print("🧪 首先进行测试模式（前3章）...")
                test_folder = crawler.crawl_book(delay=2, test_mode=True)
                
                print("\n" + "="*40)
                print("📋 测试阶段完成！")
                if test_folder:
                    print(f"📁 测试文件已保存到文件夹: {test_folder}")
                
                continue_choice = input("✨ 效果满意吗？是否继续爬取完整版本？(y/n): ").strip().lower()
                
                if continue_choice in ['y', 'yes', '是', '满意']:
                    print("\n🚀 开始完整爬取并生成完整Word文档...")
                    final_folder = crawler.crawl_book(delay=3, test_mode=False)
                    if final_folder:
                        print(f"📁 完整版文件已保存到文件夹: {final_folder}")
                else:
                    print("👋 测试完成，感谢使用！")
                break
                
            elif choice == "4":
                print("👋 感谢使用通用小说爬虫！再见！")
                break
                
            else:
                print("❌ 无效选择，请输入 1、2、3 或 4")
                
        except KeyboardInterrupt:
            print("\n\n👋 程序被用户中断，再见！")
            break
        except Exception as e:
            print(f"❌ 输入处理错误: {e}")
            print("请重新选择...")
            continue
    
    print(f"\n📁 生成的文件:")
    print(f"   📄 [书名].txt - 纯文本格式")
    print(f"   📋 [书名].json - 数据格式（包含元数据）")
    if DOCX_AVAILABLE:
        print(f"   📝 [书名].docx - Word文档（标题加粗，格式化）")
    else:
        print(f"   ⚠️  Word文档未生成（需要安装：pip install python-docx）")
    
    print(f"\n🎉 所有文件已保存在以书名命名的文件夹中！")
    print(f"📁 文件夹结构: [书名]/[书名].txt, [书名].json, [书名].docx")
    print(f"🔄 章节已按正确顺序排序，从第一章开始")
    print(f"💡 提示：文件夹和文件名根据自动识别的书名生成")

if __name__ == "__main__":
    main()