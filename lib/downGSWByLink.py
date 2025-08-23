# downList.py

import os
import urllib.request
import sys
import re
from urllib.parse import urljoin, urlparse
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
import random
from lib import downGSW

# 尝试导入 config，如果没有则使用 print 替代
try:
    import lib.config as config
except ImportError:

    class config:
        @staticmethod
        def setLog(msg):
            print(f"[LOG] {msg}")


num = 0


def progressbar(cur, total=100):
    percent = "{:.2%}".format(cur / total)
    sys.stdout.write("\r")
    sys.stdout.write("[%-100s] %s" % ("=" * int(cur), percent))
    sys.stdout.flush()


def schedule(blocknum, blocksize, totalsize):
    if totalsize == 0:
        percent = 0
    else:
        percent = blocknum * blocksize / totalsize
    if percent > 1.0:
        percent = 1.0
    percent = percent * 100
    progressbar(percent)


class down:
    def __init__(self, url):
        self.url = url
        self.base_domain = self.get_base_domain(url) if url else ""
        self.headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
            "Connection": "keep-alive",
        }
        self.book_info = {"title": "未知书名", "desc": "简介：无"}
        self.temp_path = os.path.join("下载", "temp_book.html")
        print("📘 开始下载书籍目录...")
        self.downFile()

    def get_base_domain(self, url):
        parsed = urlparse(url)
        return f"{parsed.scheme}://{parsed.netloc}"

    def downFile(self):
        # 创建目录
        os.makedirs("下载", exist_ok=True)

        # 下载主页
        opener = urllib.request.build_opener()
        opener.addheaders = [("User-Agent", self.headers["User-Agent"])]
        urllib.request.install_opener(opener)
        try:
            urllib.request.urlretrieve(self.url, self.temp_path, schedule)
        except Exception as e:
            print(f"❌ 下载失败: {e}")
            return

        # 解析章节
        articles = self.downList()
        
        if not articles:
            config.setLog("❌ 未解析到章节列表")
            return

        print(f"\n🔍 共找到 {len(articles)} 章，开始抓取...")

        # 清空旧内容
        downGSW.clear_content()

        # 逐章抓取
        for idx, item in enumerate(articles):
            print(f"[{idx+1:2d}/{len(articles)}] {item['title']}")
            downGSW.down(item["title"], item["link"])
            time.sleep(random.uniform(2.0, 4.0))  # 防爬延迟 2-4s之间

        # 合并保存
        all_content = downGSW.get_all_content()
        self.save_all_to_one_docx(all_content)
         # ✅ 在这里删除临时文件（已经不需要了）
        try:
            if os.path.exists(self.temp_path):
                os.remove(self.temp_path)
                print(f"🗑️ 临时文件已删除: {self.temp_path}")
        except Exception as e:
            print(f"⚠️ 删除临时文件失败: {e}")

    def parse_js_array(self, file_content):
        html = etree.HTML(file_content)
        a_nodes = html.xpath('//a[contains(@href,"bookv_")]')
        chapters = []
        for a in a_nodes:
            href = a.get("href")
            title = a.text.strip() if a.text else ""
            if title and "guwen" in href:
                full_url = urljoin("https://www.gushiwen.cn", href)
                chapters.append({"title": title, "link": full_url})
        return chapters

    def downList(self):
        try:
            with open(self.temp_path, "r", encoding="utf-8") as f:
                content = f.read()

            # 提取书名
            tree = etree.HTML(content)
            title_node = tree.xpath('//div[@class="main3"]//h1//b/text()')
            desc_node = tree.xpath('//div[@class="main3"]//div[@class="cont"]/p//text()')
            intro = ''.join([
                    text.strip() 
                    for text in desc_node 
                    if not text.strip().startswith('►')
                ]).strip()
            if title_node:
                self.book_info["title"] = title_node[0].strip()
            if intro:
                 self.book_info["desc"] = intro.strip()
            # 解析章节
            articles = self.parse_js_array(content)

            if not articles:
                config.setLog("❌ 未找到有效章节链接")
                return None
            # articles = articles[20:21] #测试
            # articles = articles[:1] #测试
            return articles

        except Exception as e:
            config.setLog(f"❌ 解析失败: {str(e)}")
            return None

    def save_all_to_one_docx(self, chapters_data):
        """将所有章节合并保存为一个 Word 文件，满足：标题楷体，正文段后0.2行，行距1.5倍"""
        doc = Document()

        # -------------------------------
        # 设置默认样式（Normal）
        # -------------------------------
        style = doc.styles["Normal"]
        font = style.font
        font.name = "楷体"
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        font._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")

        # -------------------------------
        # 自定义“标题1”样式（标题：楷体、黑色、小二、加粗、居中）
        # -------------------------------
        try:
            heading_style = doc.styles["Heading 1"]
        except KeyError:
            from docx.enum.style import WD_STYLE_TYPE

            heading_style = doc.styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)

        # 段落格式
        heading_format = heading_style.paragraph_format
        heading_format.space_before = Pt(24)
        heading_format.space_after = Pt(24)
        heading_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 字体格式
        font = heading_style.font
        font.name = "楷体"
        font.size = Pt(18)  # 小二
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)
        font._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")

        # -------------------------------
        # 书名（使用 Heading 1）
        # -------------------------------
        p0 = doc.add_paragraph()
        p0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p0.add_run(self.book_info["title"])
        p0.line_spacing = 1.5  # 行距 1.5 倍
       
        # 简介
        pDesc = doc.add_paragraph()
        pDesc.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        pDesc.add_run(self.book_info["desc"])
        pDesc.line_spacing = 1.5  # 行距 1.5 倍
        # 添加分页符
        doc.add_page_break()
        # 添加空行

        # -------------------------------
        # 逐章写入内容
        # -------------------------------
        for idx, (combined, paragraphs) in enumerate(chapters_data):
            # 章节标题
            p_title = doc.add_paragraph(style="Heading 1")
            p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run_title = p_title.add_run(combined)

            # ✅ 设置字体为楷体（关键：必须设置 eastAsia）
            run_title.font.name = "楷体"
            run_title.font._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")

            # ✅ 设置颜色为黑色
            run_title.font.color.rgb = RGBColor(0, 0, 0)


            # 正文段落
            for para in paragraphs:
                p = doc.add_paragraph()
                p_format = p.paragraph_format
                p_format.space_before = Pt(0)  # 段前间距 0
                p_format.space_after = Pt(
                    0.2
                )  # 段后间距 ≈ 0.2 行（小四12pt，0.2*12=2.4pt，略上浮）
                p_format.line_spacing = 1.5  # 行距 1.5 倍
                p_format.first_line_indent = Cm(0.74)  # 首行缩进 2 字符

                run = p.add_run(para['content'].strip())
                run.font.name = "楷体"
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体") 
                if para['type'] == 'text':
                    run.font.size = Pt(12)
                else:
                    run.font.size = Pt(16)
                    run.bold = True

                

            # 每章后加分页
            if idx < len(chapters_data) - 1:  # idx 从 0 开始，最后一章的 idx = len-1
                doc.add_page_break()

        # -------------------------------
        # 保存文件
        # -------------------------------
        safe_title = re.sub(r'[<>:"/\\|?*\x00]', "", self.book_info["title"])
        output_dir = os.path.join("下载")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"{safe_title}.docx")

        try:
            doc.save(output_path)
            print(f"\n🎉 成功保存：{output_path}")
            config.setLog(
                f"【{safe_title}】全本下载完成，共 {len(chapters_data)} 章，已按格式排版"
            )
        except Exception as e:
            print(f"❌ 保存失败: {e}")
            config.setLog(f"保存失败: {str(e)}")


# =============================
# 使用示例（取消注释并填写网址即可运行）
# =============================
# if __name__ == "__main__":
#     url = "https://www.gushiwen.cn/guwen/bookv_XXXXXXXX.aspx"
#     downloader = down(url)
